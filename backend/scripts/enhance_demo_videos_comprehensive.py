#!/usr/bin/env python3
"""
Comprehensive script to enhance demo video documents:
1. Analyze and improve filenames
2. Extract and add tags/categories
3. Enhance document content with metadata
4. Create improved versions with tags
"""

import os
import sys
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple

sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Install it with: pip install python-docx")
    sys.exit(1)


# Category mappings for consistent tagging
CATEGORY_MAPPINGS = {
    'dc_edge': ['DC Edge', 'Data Center Edge', 'DC', 'Edge'],
    'cloud_edge': ['Cloud Edge', 'Cloud', 'Edge'],
    'zone_seg': ['Zone Segmentation', 'Zone Seg', 'Segmentation', 'Zone'],
    'macro_micro_seg': ['Macro Micro Segmentation', 'Macro Segmentation', 'Micro Segmentation', 'Segmentation'],
    'smart_switch': ['Smart Switch', 'L4 Switch', 'Hypershield', 'Switch'],
    'ai_model_protection': ['AI Model Protection', 'AI Protection', 'AI', 'Model Protection'],
}

# Product/Feature mappings
PRODUCT_MAPPINGS = {
    'snortml': ['SnortML', 'Snort ML', 'Zero Day', 'Machine Learning'],
    'eve': ['EVE', 'Encrypted Visibility Engine', 'Encrypted Visibility'],
    'aiops': ['AIOps', 'AI Ops', 'SCC', 'Security Cloud Control'],
    'rtc': ['RTC', 'Rapid Threat Containment', 'Threat Containment'],
    'sgt': ['SGT', 'Security Group Tags', 'Tags'],
    'mcd': ['MCD', 'Automated Cloud Security Orchestration'],
}


def extract_category_from_filename(filename: str) -> Tuple[str, List[str]]:
    """Extract category and tags from filename"""
    name_lower = filename.lower()
    tags = []
    category = ""
    
    # Determine category
    if 'dc' in name_lower and 'edge' in name_lower:
        category = "DC Edge"
        tags.extend(['DC Edge', 'Data Center Edge', 'Edge Security'])
    elif 'cloud' in name_lower and 'edge' in name_lower:
        category = "Cloud Edge"
        tags.extend(['Cloud Edge', 'Cloud Security', 'Edge Security'])
    elif 'zone' in name_lower and 'seg' in name_lower:
        category = "Zone Segmentation"
        tags.extend(['Zone Segmentation', 'Segmentation', 'Network Segmentation'])
    elif 'macro' in name_lower and 'micro' in name_lower:
        category = "Macro/Micro Segmentation"
        tags.extend(['Macro Segmentation', 'Micro Segmentation', 'Segmentation'])
    elif 'smart' in name_lower and 'switch' in name_lower:
        category = "Smart Switch"
        tags.extend(['Smart Switch', 'L4 Switch', 'Switch'])
    elif 'ai' in name_lower and 'model' in name_lower:
        category = "AI Model Protection"
        tags.extend(['AI Model Protection', 'AI Security', 'Model Protection'])
    
    # Extract product/feature tags
    if 'snortml' in name_lower or 'snort' in name_lower:
        tags.extend(['SnortML', 'Zero Day Threat Defense', 'Machine Learning'])
    if 'eve' in name_lower:
        tags.extend(['EVE', 'Encrypted Visibility Engine'])
    if 'aiops' in name_lower:
        tags.extend(['AIOps', 'Security Cloud Control', 'SCC'])
    if 'rtc' in name_lower:
        tags.extend(['RTC', 'Rapid Threat Containment'])
    if 'sgt' in name_lower or 'tag' in name_lower:
        tags.extend(['SGT', 'Security Group Tags', 'Tag-based Policy'])
    if 'hypershield' in name_lower:
        tags.extend(['Hypershield', 'L4 Segmentation'])
    if 'l4' in name_lower:
        tags.extend(['L4 Switch', 'Layer 4'])
    if 'mcd' in name_lower:
        tags.extend(['MCD', 'Automated Cloud Security Orchestration'])
    
    # Add demo video tag
    tags.append('Demo Video')
    
    # Remove duplicates while preserving order
    seen = set()
    unique_tags = []
    for tag in tags:
        if tag.lower() not in seen:
            seen.add(tag.lower())
            unique_tags.append(tag)
    
    return category, unique_tags


def improve_filename(filename: str) -> Optional[str]:
    """Suggest filename improvements"""
    name = Path(filename).stem
    ext = Path(filename).suffix
    
    # Check if filename follows best practices
    issues = []
    
    # Check for underscores (should use underscores, not spaces)
    if ' ' in name and '_' not in name:
        issues.append("Contains spaces - should use underscores")
    
    # Check for proper structure
    parts = name.replace('_', ' ').lower().split()
    if 'demo' not in parts and 'video' not in parts:
        issues.append("Missing 'demo-video' suffix")
    
    # Check length
    if len(name) > 100:
        issues.append("Filename too long")
    
    if issues:
        # Suggest improvement
        improved = name
        # Ensure demo-video suffix
        if 'demo' not in parts and 'video' not in parts:
            improved = f"{improved}_demo-video"
        
        return f"{improved}{ext}"
    
    return None  # No improvement needed


def extract_youtube_links(text: str) -> List[Dict[str, str]]:
    """Extract YouTube links from text"""
    patterns = [
        r'https?://(?:www\.)?youtube\.com/watch\?v=([a-zA-Z0-9_-]+)',
        r'https?://(?:www\.)?youtu\.be/([a-zA-Z0-9_-]+)',
        r'https?://(?:www\.)?youtube\.com/embed/([a-zA-Z0-9_-]+)',
        r'youtube\.com/watch\?v=([a-zA-Z0-9_-]+)',
        r'youtu\.be/([a-zA-Z0-9_-]+)',
    ]
    
    links = []
    seen_ids = set()
    
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            video_id = match.group(1)
            if video_id not in seen_ids:
                seen_ids.add(video_id)
                url = f"https://www.youtube.com/watch?v={video_id}"
                links.append({"video_id": video_id, "url": url})
    
    return links


def extract_title_from_filename(filename: str) -> str:
    """Extract a readable title from filename"""
    name = Path(filename).stem
    # Replace underscores with spaces
    name = name.replace('_', ' ')
    # Remove 'demo-video' suffix
    name = re.sub(r'\s*demo\s*video\s*$', '', name, flags=re.IGNORECASE)
    # Clean up multiple spaces
    name = ' '.join(name.split())
    return name.strip()


def enhance_document_content(original_text: str, filename: str, category: str, tags: List[str]) -> str:
    """Enhance document content with proper structure and tags"""
    lines = original_text.split('\n')
    youtube_links = extract_youtube_links(original_text)
    title = extract_title_from_filename(filename)
    
    enhanced_lines = []
    
    # Title with category
    if category:
        enhanced_lines.append(f"{category} | {title}")
    else:
        enhanced_lines.append(title)
    enhanced_lines.append("")
    
    # Add tags as metadata (for RAG to find)
    if tags:
        enhanced_lines.append("TAGS: " + ", ".join(tags))
        enhanced_lines.append("")
    
    # Overview section
    enhanced_lines.append("OVERVIEW")
    enhanced_lines.append("")
    
    # Try to preserve existing overview
    overview_content = []
    in_overview = False
    for i, line in enumerate(lines):
        if any(keyword in line.lower() for keyword in ['overview', 'introduction']):
            in_overview = True
            continue
        if in_overview and line.strip():
            if any(keyword in line.lower() for keyword in ['key features', 'features', 'demo video', 'summary']):
                break
            overview_content.append(line)
    
    if overview_content:
        enhanced_lines.extend(overview_content)
    else:
        # Create overview with category and tags
        enhanced_lines.append(f"This document provides comprehensive information about {title}.")
        if category:
            enhanced_lines.append(f"It is part of the {category} solution portfolio.")
        enhanced_lines.append("")
    
    enhanced_lines.append("")
    enhanced_lines.append("KEY FEATURES")
    enhanced_lines.append("")
    
    # Try to preserve existing features
    features_content = []
    in_features = False
    for i, line in enumerate(lines):
        if any(keyword in line.lower() for keyword in ['key features', 'features', 'capabilities']):
            in_features = True
            continue
        if in_features and line.strip():
            if any(keyword in line.lower() for keyword in ['demo video', 'summary', 'overview']):
                break
            features_content.append(line)
    
    if features_content:
        enhanced_lines.extend(features_content)
    else:
        enhanced_lines.append("Key features and capabilities will be detailed here.")
        enhanced_lines.append("")
    
    enhanced_lines.append("")
    enhanced_lines.append("DEMO VIDEO")
    enhanced_lines.append("")
    
    # Add video links with context
    if youtube_links:
        for link in youtube_links:
            enhanced_lines.append(f"Watch the demonstration video to see {title} in action:")
            enhanced_lines.append("")
            enhanced_lines.append(f"Demo Video: {link['url']}")
            enhanced_lines.append("")
            enhanced_lines.append("This video demonstrates:")
            enhanced_lines.append("- Key features and capabilities")
            enhanced_lines.append("- Real-world use cases and scenarios")
            enhanced_lines.append("- Integration and deployment options")
            enhanced_lines.append("- Best practices and recommendations")
            enhanced_lines.append("")
    else:
        enhanced_lines.append("Demo Video: [YouTube link will be added here]")
        enhanced_lines.append("")
    
    enhanced_lines.append("SUMMARY")
    enhanced_lines.append("")
    
    # Try to preserve existing summary
    summary_content = []
    in_summary = False
    for i, line in enumerate(lines):
        if 'summary' in line.lower() or 'conclusion' in line.lower():
            in_summary = True
            continue
        if in_summary and line.strip():
            summary_content.append(line)
    
    if summary_content:
        enhanced_lines.extend(summary_content)
    else:
        enhanced_lines.append(f"{title} provides advanced capabilities for modern network security.")
        if category:
            enhanced_lines.append(f"As part of the {category} solution, it enables organizations to enhance their security posture.")
        enhanced_lines.append("")
    
    # Add source line with tags
    if youtube_links:
        source_url = youtube_links[0]['url']
        tag_str = ", ".join(tags[:3]) if tags else ""
        enhanced_lines.append("")
        if tag_str:
            enhanced_lines.append(f"Source: {title} | Tags: {tag_str} | {source_url}")
        else:
            enhanced_lines.append(f"Source: {title} | {source_url}")
    
    return '\n'.join(enhanced_lines)


def create_enhanced_docx(original_path: Path, enhanced_content: str, output_dir: Path, tags: List[str]):
    """Create an enhanced .docx file with proper formatting"""
    doc = DocxDocument()
    
    lines = enhanced_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Section headers
        if line.isupper() and len(line) > 3 and not line.startswith('HTTP') and 'TAGS:' not in line:
            p = doc.add_paragraph(line)
            p.style = 'Heading 1'
        # Tags line
        elif line.startswith('TAGS:'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)  # Gray for metadata
        # Bullet points
        elif line.startswith('-'):
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
        # Important lines (Demo Video, Source)
        elif any(keyword in line.lower() for keyword in ['demo video:', 'source:']):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        # YouTube links
        elif 'youtube.com' in line or 'youtu.be' in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0, 102, 204)  # Blue for links
            run.bold = True
        # Regular paragraphs
        else:
            doc.add_paragraph(line)
    
    # Save
    output_path = output_dir / original_path.name
    doc.save(str(output_path))
    return output_path


def main():
    """Main function"""
    script_dir = Path(__file__).parent
    project_root = script_dir.parent.parent
    demo_videos_dir = project_root / 'public' / 'demo-videos'
    
    if not demo_videos_dir.exists():
        print(f"ERROR: Directory not found: {demo_videos_dir}")
        sys.exit(1)
    
    # Create output directory
    output_dir = demo_videos_dir / 'enhanced'
    output_dir.mkdir(exist_ok=True)
    
    print("=" * 70)
    print("COMPREHENSIVE DEMO VIDEO DOCUMENT ENHANCEMENT")
    print("=" * 70)
    print(f"Source: {demo_videos_dir}")
    print(f"Output: {output_dir}")
    print()
    
    # Process files
    docx_files = [f for f in demo_videos_dir.glob('*.docx') if not f.name.startswith('~$')]
    
    if not docx_files:
        print("No .docx files found.")
        sys.exit(1)
    
    print(f"Found {len(docx_files)} document(s) to process:\n")
    
    filename_suggestions = {}
    all_tags = {}
    
    for docx_file in docx_files:
        try:
            print(f"Processing: {docx_file.name}")
            
            # Read document
            doc = DocxDocument(str(docx_file))
            original_text = '\n'.join([para.text for para in doc.paragraphs])
            
            if not original_text.strip():
                print(f"  ⚠️  Warning: Document appears empty, skipping...")
                continue
            
            # Extract category and tags
            category, tags = extract_category_from_filename(docx_file.name)
            all_tags[docx_file.name] = tags
            
            print(f"  Category: {category or 'Not identified'}")
            print(f"  Tags: {', '.join(tags[:5])}{'...' if len(tags) > 5 else ''}")
            
            # Check filename
            suggested_name = improve_filename(docx_file.name)
            if suggested_name:
                filename_suggestions[docx_file.name] = suggested_name
                print(f"  ⚠️  Filename suggestion: {suggested_name}")
            
            # Extract YouTube links
            youtube_links = extract_youtube_links(original_text)
            if not youtube_links:
                print(f"  ⚠️  Warning: No YouTube links found")
            else:
                print(f"  ✓ Found {len(youtube_links)} YouTube link(s)")
            
            # Enhance content
            enhanced_content = enhance_document_content(original_text, docx_file.name, category, tags)
            
            # Create enhanced document
            output_path = create_enhanced_docx(docx_file, enhanced_content, output_dir, tags)
            
            print(f"  ✅ Created: {output_path.name}")
            print()
            
        except Exception as e:
            print(f"  ❌ Error: {e}")
            import traceback
            traceback.print_exc()
            print()
    
    # Summary
    print("=" * 70)
    print("SUMMARY")
    print("=" * 70)
    print(f"✅ Processed {len(docx_files)} document(s)")
    print(f"✅ Enhanced documents saved to: {output_dir}")
    
    if filename_suggestions:
        print(f"\n⚠️  Filename Suggestions ({len(filename_suggestions)} files):")
        for old, new in filename_suggestions.items():
            print(f"  {old}")
            print(f"    → {new}")
    
    print(f"\n📋 Next Steps:")
    print("1. Review enhanced documents in 'enhanced' folder")
    print("2. Check that tags and categories are correct")
    print("3. Manually add more detailed content if needed")
    print("4. Consider renaming files based on suggestions")
    print("5. Replace originals and re-upload to system")
    print("\n💡 Tip: Tags are now included in document content for better RAG search!")


if __name__ == '__main__':
    main()


