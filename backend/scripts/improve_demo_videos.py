#!/usr/bin/env python3
"""
Script to improve demo video documents according to the format guide.
Reads existing .docx files, improves their structure, and saves improved versions.
"""

import os
import sys
import re
from pathlib import Path
from typing import List, Dict, Optional

# Add parent directory to path to import app modules
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Install it with: pip install python-docx")
    sys.exit(1)


def extract_youtube_links(text: str) -> List[Dict[str, str]]:
    """Extract YouTube links from text"""
    patterns = [
        r'https?://(?:www\.)?youtube\.com/watch\?v=([a-zA-Z0-9_-]+)',
        r'https?://(?:www\.)?youtu\.be/([a-zA-Z0-9_-]+)',
        r'https?://(?:www\.)?youtube\.com/embed/([a-zA-Z0-9_-]+)',
        r'youtube\.com/watch\?v=([a-zA-Z0-9_-]+)',
        r'youtu\.be/([a-zA-Z0-9_-]+)',
    ]
    
    links = []
    seen_ids = set()
    
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            video_id = match.group(1)
            if video_id not in seen_ids:
                seen_ids.add(video_id)
                url = f"https://www.youtube.com/watch?v={video_id}"
                links.append({"video_id": video_id, "url": url})
    
    return links


def extract_title_from_filename(filename: str) -> str:
    """Extract a readable title from filename"""
    # Remove extension
    name = Path(filename).stem
    # Replace underscores with spaces
    name = name.replace('_', ' ')
    # Remove 'demo-video' suffix
    name = re.sub(r'\s*demo\s*video\s*$', '', name, flags=re.IGNORECASE)
    # Clean up multiple spaces
    name = ' '.join(name.split())
    return name.strip()


def improve_document_content(original_text: str, filename: str) -> str:
    """Improve document content structure"""
    lines = original_text.split('\n')
    
    # Extract YouTube links
    youtube_links = extract_youtube_links(original_text)
    
    # Extract title from filename
    title = extract_title_from_filename(filename)
    
    # Parse category, product, description from filename
    parts = Path(filename).stem.replace('_', ' ').split()
    category = ""
    product = ""
    description = ""
    
    # Try to identify category (DC Edge, Cloud Edge, Zone Seg, etc.)
    if 'DC' in parts and 'Edge' in parts:
        category = "DC Edge"
    elif 'Cloud' in parts and 'Edge' in parts:
        category = "Cloud Edge"
    elif 'Zone' in parts and 'Seg' in parts:
        category = "Zone Segmentation"
    elif 'macro' in parts and 'micro' in parts and 'seg' in parts:
        category = "Macro/Micro Segmentation"
    elif 'smart' in parts and 'switch' in parts:
        category = "Smart Switch"
    elif 'AI' in parts and 'model' in parts:
        category = "AI Model Protection"
    
    # Build improved content
    improved_lines = []
    
    # Title
    if title:
        improved_lines.append(title)
        improved_lines.append("")
    
    # Overview section
    improved_lines.append("OVERVIEW")
    improved_lines.append("")
    
    # Try to extract existing overview or create one
    overview_found = False
    for i, line in enumerate(lines):
        if any(keyword in line.lower() for keyword in ['overview', 'introduction', 'summary']):
            # Found overview section, include next few paragraphs
            improved_lines.append(line)
            improved_lines.append("")
            for j in range(i + 1, min(i + 5, len(lines))):
                if lines[j].strip() and not lines[j].strip().startswith('#'):
                    improved_lines.append(lines[j])
                elif lines[j].strip().startswith('#'):
                    break
            overview_found = True
            break
    
    if not overview_found:
        # Create a basic overview from filename
        improved_lines.append(f"This document provides information about {title}.")
        improved_lines.append("")
    
    improved_lines.append("")
    improved_lines.append("KEY FEATURES")
    improved_lines.append("")
    
    # Try to extract features or create placeholder
    features_found = False
    for i, line in enumerate(lines):
        if any(keyword in line.lower() for keyword in ['feature', 'capability', 'benefit']):
            features_found = True
            # Include feature section
            for j in range(i, min(i + 10, len(lines))):
                if lines[j].strip():
                    improved_lines.append(lines[j])
            break
    
    if not features_found:
        improved_lines.append("Key features and capabilities will be detailed here.")
        improved_lines.append("")
    
    improved_lines.append("")
    improved_lines.append("DEMO VIDEO")
    improved_lines.append("")
    
    # Add video links
    if youtube_links:
        for link in youtube_links:
            improved_lines.append(f"Watch the demonstration video:")
            improved_lines.append("")
            improved_lines.append(f"Demo Video: {link['url']}")
            improved_lines.append("")
            improved_lines.append("This video demonstrates:")
            improved_lines.append("- Key features and capabilities")
            improved_lines.append("- Real-world use cases")
            improved_lines.append("- Integration and deployment")
            improved_lines.append("")
    else:
        improved_lines.append("Demo Video: [YouTube link will be added here]")
        improved_lines.append("")
    
    improved_lines.append("SUMMARY")
    improved_lines.append("")
    
    # Try to extract summary or create one
    summary_found = False
    for i, line in enumerate(lines):
        if 'summary' in line.lower() or 'conclusion' in line.lower():
            summary_found = True
            for j in range(i, len(lines)):
                if lines[j].strip():
                    improved_lines.append(lines[j])
            break
    
    if not summary_found:
        improved_lines.append(f"{title} provides advanced capabilities for modern network security.")
        improved_lines.append("")
    
    # Add source line
    if youtube_links:
        source_url = youtube_links[0]['url']
        improved_lines.append("")
        improved_lines.append(f"Source: {title} | {source_url}")
    
    return '\n'.join(improved_lines)


def create_improved_docx(original_path: Path, improved_content: str, output_dir: Path):
    """Create an improved .docx file"""
    doc = DocxDocument()
    
    lines = improved_content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Add empty paragraph
            doc.add_paragraph()
            continue
        
        # Check if it's a section header
        if line.isupper() and len(line) > 3 and not line.startswith('HTTP'):
            # Section header
            p = doc.add_paragraph(line)
            p.style = 'Heading 1'
            current_section = line
        elif line.startswith('-'):
            # Bullet point
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
        elif any(keyword in line.lower() for keyword in ['demo video:', 'source:']):
            # Important line - make it bold
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif 'youtube.com' in line or 'youtu.be' in line:
            # YouTube link
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color for links
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save improved document
    output_path = output_dir / original_path.name
    doc.save(str(output_path))
    print(f"✅ Created improved version: {output_path.name}")


def main():
    """Main function to improve all demo video documents"""
    # Get the demo-videos directory
    script_dir = Path(__file__).parent
    project_root = script_dir.parent.parent
    demo_videos_dir = project_root / 'public' / 'demo-videos'
    
    if not demo_videos_dir.exists():
        print(f"ERROR: Directory not found: {demo_videos_dir}")
        sys.exit(1)
    
    # Create output directory for improved files
    output_dir = demo_videos_dir / 'improved'
    output_dir.mkdir(exist_ok=True)
    
    print(f"Reading demo video documents from: {demo_videos_dir}")
    print(f"Improved versions will be saved to: {output_dir}")
    print()
    
    # Process all .docx files
    docx_files = list(demo_videos_dir.glob('*.docx'))
    # Filter out temp files (starting with ~$)
    docx_files = [f for f in docx_files if not f.name.startswith('~$')]
    
    if not docx_files:
        print("No .docx files found in the directory.")
        sys.exit(1)
    
    print(f"Found {len(docx_files)} document(s) to process:\n")
    
    for docx_file in docx_files:
        try:
            print(f"Processing: {docx_file.name}")
            
            # Read original document
            doc = DocxDocument(str(docx_file))
            original_text = '\n'.join([para.text for para in doc.paragraphs])
            
            if not original_text.strip():
                print(f"  ⚠️  Warning: Document appears to be empty, skipping...")
                continue
            
            # Extract YouTube links
            youtube_links = extract_youtube_links(original_text)
            if not youtube_links:
                print(f"  ⚠️  Warning: No YouTube links found in document")
            
            # Improve content
            improved_content = improve_document_content(original_text, docx_file.name)
            
            # Create improved document
            create_improved_docx(docx_file, improved_content, output_dir)
            
            print(f"  ✓ Processed successfully")
            if youtube_links:
                print(f"    Found {len(youtube_links)} YouTube link(s)")
            print()
            
        except Exception as e:
            print(f"  ❌ Error processing {docx_file.name}: {e}")
            print()
    
    print(f"\n✅ Done! Improved documents saved to: {output_dir}")
    print("\nNext steps:")
    print("1. Review the improved documents in the 'improved' folder")
    print("2. Manually enhance them with more detailed content if needed")
    print("3. Replace the original files with improved versions")
    print("4. Re-upload to the system")


if __name__ == '__main__':
    main()


